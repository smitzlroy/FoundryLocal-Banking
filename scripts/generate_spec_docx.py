"""Generate the Sovereign Banking Rate-Forecast demo specification as a Word .docx.

Run from the repo root inside the project venv:
    python scripts/generate_spec_docx.py

Produces: docs/SPEC.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
# Palette
# ---------------------------------------------------------------------------
AZURE_BLUE = RGBColor(0x00, 0x78, 0xD4)
EDGE_GREEN = RGBColor(0x2E, 0x7D, 0x32)
DARK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x60, 0x60, 0x60)
HDR_FILL = "0078D4"
EDGE_FILL = "2E7D32"


def shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, white=False, size=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def add_table(doc, headers, rows, *, header_fill=HDR_FILL, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, white=True, size=10)
        shade_cell(hdr[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in table.rows:
                r.cells[i].width = Inches(w)
    return table


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def body(doc, text, *, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def mono(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.makeelement(qn("w:rFonts"), {})
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rpr.append(rfonts)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    return p


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
def build() -> Document:
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK

    # ---- Title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Sovereign Banking Rate-Forecast Demo")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = AZURE_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Cloud \u2192 Edge Model Lifecycle on Foundry Local (Azure Local / AKS Arc)")
    sr.font.size = Pt(15)
    sr.font.color.rgb = GREY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Technical Specification \u2014 Draft v0.1 \u2014 11 June 2026")
    mr.font.size = Pt(11)
    mr.font.color.rgb = GREY

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m2 = meta2.add_run("Cluster: fl-banking-portland  \u2022  Resource group: sovereign-ai-daz")
    m2.font.size = Pt(10)
    m2.font.color.rgb = GREY

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Goal: prove the operational lifecycle of taking a predictive model that runs in Azure "
        "today and migrating it \u2014 together with its UI \u2014 to run entirely at the edge on "
        "Azure Local, then augmenting it with an open-source catalog model."
    )
    nr.italic = True
    nr.font.size = Pt(11)
    nr.font.color.rgb = DARK

    doc.add_page_break()

    # ---- 1. Purpose & narrative ----
    h(doc, "1. Purpose & Narrative", 1)
    body(doc,
         "This demo proves the operational lifecycle of taking a predictive model that \u201Cruns in "
         "Azure today\u201D and bringing it to Azure Local at the edge via Foundry Local \u2014 for a "
         "regulated bank that needs data sovereignty, low latency, and no cloud egress. The emphasis is "
         "on the lifecycle realities (convert \u2192 package \u2192 push \u2192 deploy \u2192 manage \u2192 "
         "repoint), not on raw inference performance.")
    body(doc,
         "It mirrors the customer scenario: they run models in Azure and want to take the same model, "
         "convert it to ONNX, containerize it, and deploy it to Azure Local at the edge \u2014 then add an "
         "open-source model from the catalog for additional intelligence. The demo deliberately begins at "
         "the ONNX handoff contract: converting a model to ONNX is the customer\u2019s responsibility, and the "
         "platform value is the paved road from a validated ONNX model onward (see Section 5). Crucially, "
         "once the workload is migrated, EVERYTHING runs at the edge \u2014 the model and the user interface "
         "\u2014 with no cloud dependency in the running state.")

    # ---- 2. The predictive model ----
    h(doc, "2. The Predictive Model (the BYO Hero)", 1)
    add_table(doc,
              ["Property", "Detail"],
              [
                  ["Type", "Multi-output regression \u2014 HistGradientBoostingRegressor (scikit-learn)"],
                  ["Task", "Forecast a 9-tenor yield curve (1M, 3M, 6M, 1Y, 2Y, 3Y, 5Y, 7Y, 10Y) from a 10-day lookback (90 features)"],
                  ["Data source", "Synthetic \u2014 Vasicek short-rate + Nelson-Siegel curve (generate_data.py \u2192 rates.csv, 2,520 rows). Regulator-safe; no real customer data"],
                  ["Export", "scikit-learn \u2192 ONNX opset 17 via skl2onnx (train_export.py); validated by validate_onnx.py; test MAE 9.62 bps"],
                  ["Artifact", "model/artifacts/rate_forecast.onnx (1.4 MB)"],
              ],
              col_widths=[1.6, 4.9])
    body(doc,
         "This model stands in for the bank\u2019s existing in-Azure model. Real customers export their "
         "Azure ML or trained model to ONNX the same way (skl2onnx / torch.onnx / optimum / Olive).",
         italic=True, color=GREY)

    # ---- 3. ONNX vs vLLM ----
    h(doc, "3. ONNX vs vLLM \u2014 Why ONNX is the Spine", 1)
    body(doc,
         "A common question for BYO-at-scale is whether to standardize on vLLM. The two engines serve "
         "different purposes:")
    add_table(doc,
              ["Dimension", "ONNX Runtime (onnx / onnx-genai)", "vLLM (vllm)"],
              [
                  ["Model types", "Predictive ML, embeddings, transformers, and generative LLMs", "Generative transformer LLMs only \u2014 cannot serve predictive models"],
                  ["Hardware", "CPU, GPU, NPU, ARM, web \u2014 runs anywhere", "GPU-first (CUDA/ROCm); CPU backend exists but is not the target"],
                  ["Edge / CPU-only", "Ideal \u2014 lightweight, no GPU required", "Poor fit \u2014 designed for GPU servers"],
                  ["LLM throughput at scale", "Good single/low concurrency", "Best-in-class \u2014 PagedAttention + continuous batching"],
                  ["Best for", "Predictive ML, edge, CPU, portability, small models", "High-concurrency LLM serving on GPU"],
              ],
              col_widths=[1.5, 2.7, 2.3])
    body(doc,
         "Decision: the demo\u2019s hero spine is the predictive ONNX rate model (Acts 1\u20132). vLLM is "
         "the right tool for LLMs at scale on GPU and is presented as the documented production path \u2014 "
         "an authored manifest with a one-line runtime swap \u2014 but not run live, because the target "
         "cluster is CPU-only. The generative model in Act 3 is the catalog Phi-4 model, run live on CPU "
         "using the onnx-genai runtime.",
         color=EDGE_GREEN)

    # ---- 4. Azure cloud side (Act 1) ----
    h(doc, "4. Azure (Cloud) Side \u2014 Act 1 \u201CModel in Azure Today\u201D", 1)
    body(doc, "Act 1 establishes the \u201Cbefore\u201D baseline that we subsequently migrate away from.")
    bullet(doc, "hosts a small FastAPI + onnxruntime service exposing /v1/predict (the dashboard contract).", bold_lead="Azure Container Apps ")
    bullet(doc, "(flbankingacr, Standard) stores the inference image and, later, the model OCI artifact.", bold_lead="Azure Container Registry ")
    bullet(doc, "(\u201CSovereign Rate Forecast\u201D, Next.js) calls the Container Apps endpoint.", bold_lead="Dashboard UI ")
    bullet(doc, "(SPN fl-banking-foundry-auth) for registry pull; optional Log Analytics / App Insights.", bold_lead="Microsoft Entra ID ")
    body(doc,
         "Note: this is Azure Container Apps in the PUBLIC CLOUD (managed, serverless) \u2014 distinct from "
         "\u201CContainer Apps on Arc.\u201D See Section 7.", italic=True, color=GREY)

    # ---- 5. Scope boundary + packaging ----
    h(doc, "5. Scope Boundary, Containerize & Push (the Lifecycle Mechanics)", 1)
    body(doc,
         "The demo deliberately STARTS from a validated ONNX model. Converting a model to ONNX is the "
         "customer\u2019s responsibility: it is model-specific and involves customer-owned decisions "
         "(opset, quantization, accuracy validation, tokenizer handling). It is not Microsoft\u2019s role to "
         "convert customers\u2019 models for them \u2014 they research, test, and plan that against their own "
         "models. Microsoft\u2019s value is everything downstream of ONNX: the paved road from a validated "
         "ONNX model to a managed, sovereign edge deployment. \u201CWhen you get to ONNX, here is what to do.\u201D",
         color=EDGE_GREEN)
    body(doc, "Responsibility boundary (RACI):", space_after=2)
    add_table(doc,
              ["Stage", "Owner"],
              [
                  ["Model selection, training, conversion to ONNX, accuracy validation", "Customer"],
                  ["ONNX \u2192 package \u2192 ACR \u2192 deploy \u2192 serve \u2192 manage at the edge", "Microsoft platform (Foundry Local)"],
              ],
              col_widths=[4.6, 1.9])
    body(doc, "Lifecycle steps shown in the demo (all downstream of ONNX):", space_after=2)
    bullet(doc, "Start from a validated ONNX model (the BYO handoff contract).")
    bullet(doc, "Packaging for Foundry Local: tar.gz the ONNX + metadata, then push to ACR as an OCI artifact via ORAS (media-type application/vnd.foundrylocal.model.v1+tar) using package_and_push.sh. This is the \u201Ccontainerize the model\u201D step.")
    bullet(doc, "Deploy to Foundry Local on the edge cluster, serve, and manage.")
    mono(doc, "./model/package_and_push.sh ./artifacts flbankingacr models/rate-forecast v1")
    body(doc,
         "Getting to ONNX (guidance only \u2014 not in the live demo path): customers convert their own "
         "models using the appropriate path \u2014 skl2onnx (scikit-learn), torch.onnx (PyTorch), optimum or "
         "Olive (transformers / LLMs) \u2014 followed by accuracy validation against the source model. The "
         "convert_hf_to_onnx.sh script is kept in the repository as an illustrative, pre-baked example "
         "(\u201Ccustomers do this for their own models\u201D); it is not executed during the demo.",
         italic=True, color=GREY)

    # ---- 6. Azure Local edge side ----
    h(doc, "6. Azure Local (Edge) Side \u2014 Acts 2/3 \u201CEverything at the Edge\u201D", 1)
    body(doc, "Target cluster (deployed and verified):")
    add_table(doc,
              ["Component", "Detail"],
              [
                  ["Cluster", "fl-banking-portland (AKS Arc on Azure Local)"],
                  ["Nodes", "2 \u00d7 Standard_D8s_v3 (8 vCPU / 32 GB), CPU-only"],
                  ["Platform", "Kubernetes 1.33.5, Azure Linux 3.0, calico; custom location Portland; LNET pdx-lnet-vlan32"],
                  ["Foundry Local", "Arc Kubernetes extension Microsoft.Foundry (the \u201Cinference operator\u201D)"],
                  ["Operator pods", "api, inference-operator, model-store, telemetry-collector (ns foundry-local-operator)"],
                  ["CRDs", "modeldeployments, inferenceservices, models, storemodels (.foundrylocal.azure.com)"],
                  ["Runtimes", "onnx (predictive), onnx-genai (generative CPU), vllm (generative GPU)"],
                  ["Catalog", "~179 curated open-source models (e.g. Phi-4-generic-cpu)"],
                  ["Prereqs installed", "cert-manager (azure-cert-manager), ingress-nginx (Helm, LoadBalancer)"],
              ],
              col_widths=[1.7, 4.8], header_fill=EDGE_FILL)
    body(doc, "Deployments on the cluster:")
    bullet(doc, "BYO predictive \u2014 modeldeployment-rate-forecast.yaml: model.custom (ACR registry/repo/tag + credentials.secretRef=registry-credentials), workloadType predictive, runtime onnx, compute cpu.")
    bullet(doc, "Catalog LLM \u2014 modeldeployment-catalog-smoke.yaml: model.catalog.name Phi-4-generic-cpu, runtime onnx-genai, workloadType generative, compute cpu.")
    bullet(doc, "Dashboard UI \u2014 a plain Kubernetes Deployment + Service + Ingress on the SAME cluster, with FOUNDRY_ENDPOINT pointing at the in-cluster model service.")
    body(doc,
         "Endpoints via ingress: predictive /v1/predict; generative /v1/chat/completions. API key in the "
         "<deployment>-api-keys secret. Operator access from the laptop is via az connectedk8s proxy.")

    # ---- 7. Container Apps on Arc ----
    h(doc, "7. Do We Need Container Apps on Arc? \u2014 No", 1)
    body(doc,
         "Directive: once migrated to the edge, everything runs at the edge on fl-banking-portland, "
         "including the UI, with no cloud dependency in the running state.", color=EDGE_GREEN)
    bullet(doc, "The edge needs only: Foundry Local (model serving) + cert-manager + ingress-nginx. The dashboard UI runs as a plain Kubernetes Deployment + Service + Ingress on the same AKS Arc cluster \u2014 no additional extension.")
    bullet(doc, "Container Apps on Arc is rejected for this demo: it adds a second extension and, on AKS Arc, is Linux-only and requires a LoadBalancer plus custom CoreDNS and HAProxy, has no managed identities (an app service principal must be used to pull from ACR), is region-restricted, and sends logs to a per-cluster Log Analytics workspace. It is mentioned in the talk track only.")
    bullet(doc, "On-stage distinction: Act 1 uses Azure Container Apps in the PUBLIC CLOUD (the \u201Cbefore\u201D baseline). After migration, the model AND the UI are 100% on Azure Local \u2014 fully sovereign and able to run air-gapped. ACR is the one cloud touchpoint during deployment (artifact + image pull); once cached, the running state is sovereign.")

    # ---- 8. End-to-end flow ----
    h(doc, "8. End-to-End Flow (per Act)", 1)
    add_table(doc,
              ["Act", "What happens", "Where"],
              [
                  ["Act 1 \u2014 Cloud baseline", "Synthetic data \u2192 ONNX model \u2192 FastAPI inference \u2192 dashboard live", "Azure Container Apps (public cloud)"],
                  ["Act 2 \u2014 Edge migration (HERO)", "Same rate_forecast.onnx \u2192 ORAS package \u2192 ACR \u2192 ModelDeployment (onnx, cpu) on Foundry Local; UI redeployed to the cluster \u2192 identical curve, now sovereign", "Azure Local \u2014 fl-banking-portland"],
                  ["Act 3 \u2014 Add intelligence", "Catalog Phi-4 (onnx-genai) \u2192 dashboard \u201CAI Analyst\u201D panel; authored vLLM GPU manifest (documented, not run). Conversion shown only as an illustrative artifact", "Azure Local \u2014 fl-banking-portland"],
              ],
              col_widths=[1.7, 3.6, 1.2])

    # ---- 9. Azure services summary ----
    h(doc, "9. Azure Services Summary", 1)
    body(doc, "Cloud (Act 1 baseline):", space_after=2)
    bullet(doc, "Azure Container Apps \u2014 Act 1 inference host")
    bullet(doc, "Azure Container Registry (flbankingacr) \u2014 images and model OCI artifacts")
    bullet(doc, "Microsoft Entra ID \u2014 app registration / SPN for ACR pull")
    bullet(doc, "Optional: Log Analytics / Application Insights")
    body(doc, "Edge (Azure Local):", space_after=2)
    bullet(doc, "Azure Local + AKS Arc (Arc-enabled Kubernetes)")
    bullet(doc, "Arc Resource Bridge, Custom Location")
    bullet(doc, "Foundry Local extension (Microsoft.Foundry)")
    bullet(doc, "cert-manager extension, ingress-nginx")
    body(doc, "All edge resources are managed centrally from Azure via Arc.", italic=True, color=GREY)

    # ---- 10. Build deliverables ----
    h(doc, "10. Build Deliverables", 1)
    add_table(doc,
              ["Deliverable", "Purpose / Act"],
              [
                  ["azure-api/ (FastAPI + onnxruntime + Dockerfile)", "Act 1 cloud inference service"],
                  ["infra/20-deploy-aca.ps1", "Act 1 \u2014 build + deploy to Azure Container Apps"],
                  ["k8s/modeldeployment-rate-forecast.yaml (filled) + registry secret", "Act 2 \u2014 BYO predictive model on the edge"],
                  ["k8s/dashboard-deployment.yaml (Deployment + Service + Ingress)", "Act 2 \u2014 UI running entirely at the edge"],
                  ["app/app/api/insight/route.ts + analyst panel component", "Act 3 \u2014 LLM \u201CAI Analyst\u201D layer (catalog Phi-4)"],
                  ["k8s/modeldeployment-llm-vllm.yaml", "Act 3 \u2014 documented GPU production path (not run)"],
                  ["model/convert_hf_to_onnx.sh (illustrative)", "Reference only \u2014 \u201Cgetting to ONNX\u201D example, not run in demo"],
                  ["docs/DEMO.md", "Workshop runbook"],
              ],
              col_widths=[3.6, 2.9])

    # ---- 11. Risks & safety ----
    h(doc, "11. Risks & Safety", 1)
    bullet(doc, "CPU-only \u2192 slow generative latency. Mitigation: small models + \u201Cprocess not performance\u201D framing.")
    bullet(doc, "Network drop at the workshop. Mitigation: dashboard mock fallbacks per act.")
    bullet(doc, "vLLM cannot run (no GPU). Mitigation: manifest authored and validated with kubectl --dry-run=server only.")
    bullet(doc, "Edge access only via az connectedk8s proxy. Mitigation: keep the proxy alive during the demo.")
    bullet(doc, "Operational guardrail: stay in resource group sovereign-ai-daz; only create/edit/delete fl-banking-* resources; everything else is read-only.")

    # ---- 12. Open decisions ----
    h(doc, "12. Open Decisions", 1)
    add_table(doc,
              ["#", "Decision", "Status"],
              [
                  ["1", "Edge UI host = plain Kubernetes Deployment on fl-banking-portland; everything at the edge after migration", "Resolved"],
                  ["2", "Act 3 uses the catalog Phi-4 model only; no live conversion. Conversion to ONNX is the customer\u2019s responsibility, shown via an illustrative script only", "Resolved"],
                  ["3", "Spec output format = real Word .docx (docs/SPEC.docx)", "Resolved"],
              ],
              col_widths=[0.4, 4.9, 1.2])

    return doc


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_dir = repo_root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "SPEC.docx"
    doc = build()
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
